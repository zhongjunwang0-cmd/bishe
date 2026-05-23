# -*- coding: utf-8 -*-
"""Generate Chapter 4 Word outline (.docx and .doc if Word is available)."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS = Path(__file__).resolve().parents[1]
OUT_DOCX = DOCS / "第四章 详细设计与实现.docx"
OUT_DOC = DOCS / "第四章 详细设计与实现.doc"


def set_run_font(run, name: str = "宋体", size: int = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, "黑体", {1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    shading = p._element.get_or_add_pPr()
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10.5)
    doc.add_paragraph()
    return table


def add_figure_placeholder(doc: Document, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【此处插入截图：{caption}】")
    set_run_font(run, size=10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=10.5)


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("第四章  详细设计与实现")
    set_run_font(tr, "黑体", 18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Brix 英语学习系统")
    set_run_font(sr, "宋体", 14)

    note = doc.add_paragraph()
    nr = note.add_run(
        "说明：灰色占位处请插入系统运行截图；流程图可从 docs/figures/fig-4-*.md 用 Mermaid 导出 PNG 后粘贴。"
    )
    set_run_font(nr, size=10)
    nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # 4.1
    add_heading(doc, "4.1  系统总体设计", 1)
    add_heading(doc, "4.1.1  设计目标与原则", 2)
    add_body(
        doc,
        "在第三章需求分析的基础上，本章从可实现的工程视角阐述 Brix 英语学习系统的详细设计。"
        "设计遵循前后端分离、领域分层、智能服务解耦与数据驱动个性化四项原则。"
        "系统逻辑架构见第三章图 3-1，本章重点说明各模块内部结构、关键算法、接口契约与界面实现。",
    )
    add_bullet(doc, "前后端分离：Vue 3 与 Spring Boot 通过 REST/JSON 通信。")
    add_bullet(doc, "领域分层：Controller → Service → Mapper，统一 Result<T> 响应。")
    add_bullet(doc, "智能服务解耦：ai-service（FastAPI）独立部署深度学习推理。")
    add_bullet(doc, "数据驱动：MySQL 持久化 + Redis 缓存高频只读列表。")

    add_heading(doc, "4.1.2  工程目录与分层结构", 2)
    add_table(
        doc,
        ["层次", "技术栈", "主要目录", "职责"],
        [
            ["表现层", "Vue 3 + Element Plus", "frontend/src/views", "页面交互、路由、API 调用"],
            ["业务层", "Spring Boot 3 + Shiro", "backend/.../controller、service", "鉴权、业务逻辑"],
            ["数据层", "MySQL 8、Redis 7", "mysql/init", "表结构、迁移脚本"],
            ["智能层", "FastAPI + PyTorch", "ai-service/app、ml/models", "DKT、T5、Whisper"],
        ],
    )

    add_heading(doc, "4.1.3  统一接口规范", 2)
    add_body(doc, "所有业务 API 前缀为 /api，响应体结构如下：")
    add_code(
        doc,
        "public class Result<T> {\n"
        "    private Integer code;   // 200 成功\n"
        "    private String message;\n"
        "    private T data;\n"
        "}",
    )
    add_body(
        doc,
        "Shiro 对 /api/** 做会话校验；管理端接口要求 ADMIN 或 TEACHER 角色。",
    )

    # 4.2
    add_heading(doc, "4.2  数据库详细设计", 1)
    add_heading(doc, "4.2.1  概念模型", 2)
    add_bullet(doc, "用户域：sys_user、sys_role、sys_permission（RBAC）。")
    add_bullet(doc, "内容域：biz_vocab、biz_grammar、biz_question_bank 及阅读/听力/填空实例表。")
    add_bullet(doc, "行为域：biz_user_vocab、biz_learning_record、biz_learning_plan_item、biz_discuss。")

    add_heading(doc, "4.2.2  核心表说明", 2)
    add_body(doc, "（1）sys_user：存储账号、BCrypt 密码、level、target_exam、daily_goal 等用户画像。")
    add_body(doc, "（2）biz_user_vocab：生词本与艾宾浩斯复习状态。")
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["user_id, vocab_id", "BIGINT", "联合唯一"],
            ["mastery_level", "INT", "掌握度 0–100"],
            ["review_stage", "INT", "复习阶段 0–6"],
            ["status", "VARCHAR", "NEW/LEARNING/REVIEWING/MASTERED"],
            ["next_review_time", "DATETIME", "下次复习时间"],
        ],
    )
    add_body(
        doc,
        "（3）biz_question_bank：教师维护题库，questions_json/answers_json 存题目与答案；"
        "学习者 /generate 随机抽题生成测验实例。",
    )
    add_body(doc, "（4）biz_learning_record：各模块学习时长与得分，供统计与 DKT 推荐使用。")
    add_body(doc, "（5）biz_learning_plan_item：周计划任务，source 区分 template 与 ai_recommend。")

    add_heading(doc, "4.2.3  索引与完整性", 2)
    add_bullet(doc, "生词本 uk_user_vocab(user_id, vocab_id) 防重复收藏。")
    add_bullet(doc, "周计划 idx_user_week、idx_user_date 加速查询。")

    # 4.3
    add_heading(doc, "4.3  后端模块设计与实现", 1)
    add_heading(doc, "4.3.1  认证与权限模块", 2)
    add_body(doc, "采用 Apache Shiro，UserRealm 完成认证；密码 BCrypt 比对。按 role_id 映射 ADMIN/TEACHER/USER。")
    add_code(
        doc,
        "if (user.getRoleId() == 1L) info.addRole(\"ADMIN\");\n"
        "else if (user.getRoleId() == 2L) info.addRole(\"TEACHER\");\n"
        "else info.addRole(\"USER\");",
    )
    add_figure_placeholder(doc, "图 4-1  用户登录界面")

    add_heading(doc, "4.3.2  词汇与生词本模块", 2)
    add_table(
        doc,
        ["接口", "方法", "说明"],
        [
            ["/api/vocab/book/list", "GET", "生词本列表"],
            ["/api/vocab/book/add/{vocabId}", "POST", "加入生词本"],
            ["/api/vocab/review/today", "GET", "今日待复习"],
            ["/api/vocab/review/submit", "POST", "提交复习结果"],
        ],
    )
    add_body(
        doc,
        "艾宾浩斯算法（EbbinghausUtil）：复习间隔 1→2→4→7→15→30 天；"
        "记得则阶段+1、掌握度+15；忘记则阶段回退、掌握度-20；掌握度≥85 且阶段≥4 为 MASTERED。",
    )
    add_code(
        doc,
        "private static final int[] INTERVAL_DAYS = {1, 2, 4, 7, 15, 30};\n"
        "public static LocalDateTime nextReviewTime(int stage) {\n"
        "    return LocalDateTime.now().plusDays(intervalDaysForStage(stage));\n"
        "}",
    )

    add_heading(doc, "4.3.3  测评类模块", 2)
    add_body(
        doc,
        "阅读/听力/填空共享「抽题—作答—计分—写记录」流程："
        "POST .../generate → GET .../{id} → POST .../{id}/submit（ScoreUtil 计分）。",
    )

    add_heading(doc, "4.3.4  学习记录与统计", 2)
    add_body(doc, "RecordController、StatsController 聚合练习数据，供首页 ECharts 与 DKT 模块统计使用。")

    add_heading(doc, "4.3.5  AI 代理模块", 2)
    add_bullet(doc, "薄弱模块推荐：近 30 天记录 → ai-service DKT → 同步周计划。")
    add_bullet(doc, "语法纠错：T5-GEC，失败回退规则引擎。")
    add_bullet(doc, "AI 问答：sys_config 配置 LLM Key，未配置则用规则回复。")

    add_heading(doc, "4.3.6  管理后台", 2)
    add_body(doc, "Admin*Controller 提供题库、用户、系统配置维护；写操作 @CacheEvict 失效 Redis。")

    # 4.4
    add_heading(doc, "4.4  前端界面设计与实现", 1)
    add_heading(doc, "4.4.1  路由与布局", 2)
    add_table(
        doc,
        ["路由", "组件", "功能"],
        [
            ["/", "Home.vue", "首页看板、周计划"],
            ["/vocab", "VocabLearning.vue", "词库、生词本、闪卡复习"],
            ["/reading", "ReadingPractice.vue", "阅读测验"],
            ["/listening", "ListeningTraining.vue", "听力训练"],
            ["/oral", "OralPractice.vue", "口语录音评测"],
            ["/ai-tutoring", "AiTutoring.vue", "AI 辅导"],
            ["/admin/*", "AdminDashboard 等", "管理后台"],
        ],
    )

    add_heading(doc, "4.4.2  词汇学习界面", 2)
    add_body(doc, "三标签页：词库浏览、我的生词本（掌握度进度条）、今日复习（闪卡记得/忘记）。")
    add_figure_placeholder(doc, "图 4-2  词汇词库浏览界面")
    add_figure_placeholder(doc, "图 4-3  词汇闪卡复习界面")

    add_heading(doc, "4.4.3  首页与学习看板", 2)
    add_figure_placeholder(doc, "图 4-4  学习者首页看板")

    add_heading(doc, "4.4.4  测评类界面", 2)
    add_figure_placeholder(doc, "图 4-5  阅读理解测验界面")
    add_figure_placeholder(doc, "图 4-6  口语发音评测结果")

    add_heading(doc, "4.4.5  个人中心与 AI 辅导", 2)
    add_figure_placeholder(doc, "图 4-7  AI 辅导界面")

    add_heading(doc, "4.4.6  管理后台界面", 2)
    add_figure_placeholder(doc, "图 4-8  题库管理界面")

    # 4.5
    add_heading(doc, "4.5  智能推理服务设计与实现", 1)
    add_table(
        doc,
        ["端点", "模型", "功能"],
        [
            ["POST /api/kt/recommend", "LSTM-DKT", "薄弱模块推荐"],
            ["POST /api/grammar/correct", "T5-GEC", "语法纠错"],
            ["POST /api/pronunciation/score", "Whisper-WER", "发音评测"],
        ],
    )
    add_body(doc, "推荐优先级：DKT → GBDT 备用 → 规则。Java 通过 ai.service.base-url 调用。")

    # 4.6
    add_heading(doc, "4.6  关键业务流程", 1)
    add_heading(doc, "4.6.1  词汇复习流程", 2)
    add_bullet(doc, "GET /api/vocab/review/today 获取到期单词。")
    add_bullet(doc, "POST /api/vocab/review/submit 更新艾宾浩斯状态。")
    add_figure_placeholder(doc, "图 4-9  词汇复习流程序列图（见 docs/figures/fig-4-1）")

    add_heading(doc, "4.6.2  阅读测验流程", 2)
    add_bullet(doc, "generate → 答题 → submit 计分 → 写入 learning_record。")
    add_figure_placeholder(doc, "图 4-10  阅读测验流程序列图（见 docs/figures/fig-4-2）")

    add_heading(doc, "4.6.3  AI 个性化推荐流程", 2)
    add_bullet(doc, "聚合记录 → DKT 推荐 → syncAiRecommendToPlan → 首页/个人中心展示。")

    # 4.7 4.8
    add_heading(doc, "4.7  部署与运行实现", 1)
    add_body(
        doc,
        "Docker Compose 启动 MySQL、Redis、backend、frontend（Nginx）、ai-service。"
        "详见项目 docs/DEPLOYMENT.md。",
    )

    add_heading(doc, "4.8  本章小结", 1)
    add_body(
        doc,
        "本章完成了数据库、后端模块、前端界面及 AI 服务的详细设计与实现说明，"
        "并给出词汇复习、阅读测验、智能推荐等核心流程。"
        "请补充 6–8 张界面截图与 2–3 张流程图，全章建议不少于 5 页。",
    )

    add_heading(doc, "附录  图表清单", 1)
    add_table(
        doc,
        ["编号", "图题", "来源"],
        [
            ["图 4-1", "用户登录界面", "Login.vue 截图"],
            ["图 4-2", "词汇词库浏览", "VocabLearning 浏览 Tab"],
            ["图 4-3", "词汇闪卡复习", "VocabLearning 复习 Tab"],
            ["图 4-4", "学习者首页", "Home.vue"],
            ["图 4-5", "阅读测验", "ReadingPractice"],
            ["图 4-6", "口语评测", "OralPractice"],
            ["图 4-7", "AI 辅导", "AiTutoring"],
            ["图 4-8", "题库管理", "TestManagement"],
            ["图 4-9", "词汇复习序列图", "fig-4-1-vocab-review-sequence.md"],
            ["图 4-10", "阅读测验序列图", "fig-4-2-reading-exam-sequence.md"],
        ],
    )

    return doc


def convert_docx_to_doc(docx_path: Path, doc_path: Path) -> bool:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 0 = wdFormatDocument (.doc)
        doc.SaveAs(str(doc_path.resolve()), FileFormat=0)
        doc.Close()
        return True
    except Exception as exc:
        print(f"Word COM conversion skipped: {exc}", file=sys.stderr)
        return False
    finally:
        if word is not None:
            word.Quit()


def main():
    doc = build_document()
    doc.save(OUT_DOCX)
    print(f"Created: {OUT_DOCX}")

    if convert_docx_to_doc(OUT_DOCX, OUT_DOC):
        print(f"Created: {OUT_DOC}")
    else:
        # Fallback: Word 2003 XML wrapped as .doc (opens in Word/WPS)
        html_doc = DOCS / "第四章 详细设计与实现.doc"
        if not OUT_DOC.exists():
            print(
                "Note: .doc not created (Word/pywin32 unavailable). "
                "Open the .docx file directly, or install pywin32 for .doc export.",
                file=sys.stderr,
            )


if __name__ == "__main__":
    main()
